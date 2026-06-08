from __future__ import annotations

import argparse
import os
import re
from dataclasses import asdict, dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from Altomizer.preserver import DocumentPreserver

PREFIX_RE = re.compile(
    r"^(?P<leading>[\s\t]*)(?P<opener>[\(\[]?)(?P<value>[0-9]+|[a-zA-Z]+)(?P<closer>[\.\)\]])(?P<spacing>[\s\t]+)"
)
BULLET_RE = re.compile(
    r"^(?P<leading>[\s\t]*)(?P<marker>[\u2022\u25CF\u25CB\u25AA\u25FE\u25B6\u2043\u2219\u2013\u2014\-\*\uf000-\uf0ff])(?P<spacing>[\s\t]+)"
)
STYLE_LEVEL_RE = re.compile(r"^List (?:Bullet|Number)\s+(\d+)$", re.IGNORECASE)
BULLET_SYMBOLS = ("\u25CF", "\u25CB", "\u25AA")

PPR_ORDER = {
    tag: index
    for index, tag in enumerate(
        [
            "pStyle",
            "keepNext",
            "keepLines",
            "pageBreakBefore",
            "framePr",
            "widowControl",
            "numPr",
            "pBdr",
            "shd",
            "tabs",
            "spacing",
            "ind",
            "contextualSpacing",
            "jc",
            "outlineLvl",
            "rPr",
            "sectPr",
        ]
    )
}


@dataclass
class ListFixStats:
    ordered_tagged: int = 0
    unordered_tagged: int = 0
    bullet_templates_standardized: int = 0

    def to_dict(self) -> dict[str, int]:
        return asdict(self)


def oxpath(el, query):
    return etree._Element.xpath(el, query, namespaces=el.nsmap if hasattr(el, "nsmap") else None)


def has_numpr(paragraph) -> bool:
    p_pr = paragraph._element.find(qn("w:pPr"))
    return p_pr is not None and p_pr.find(qn("w:numPr")) is not None


def build_numPr(ilvl: int, num_id: int) -> OxmlElement:
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    return num_pr


def _insert_pPr_child(p_pr, new_el) -> None:
    tag_name = new_el.tag.split("}")[-1]
    target_idx = PPR_ORDER.get(tag_name, 99)

    existing = p_pr.find(qn(f"w:{tag_name}"))
    if existing is not None:
        p_pr.remove(existing)

    for child in p_pr:
        child_tag = child.tag.split("}")[-1]
        if PPR_ORDER.get(child_tag, 99) > target_idx:
            child.addprevious(new_el)
            return
    p_pr.append(new_el)


def clear_paragraph_indents(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.left_indent = None
    fmt.first_line_indent = None


def is_roman(value: str) -> bool:
    normalized = value.strip(" ().[]").upper()
    if not normalized:
        return False
    if len(normalized) == 1 and normalized not in ("I", "V", "X"):
        return False
    roman_chars = set("IVXLCDM")
    if not all(char in roman_chars for char in normalized):
        return False
    return bool(re.match(r"^M{0,3}(CM|CD|D?C{0,3})(XC|XL|L?X{0,3})(IX|IV|V?I{0,3})$", normalized))


def roman_to_int(value: str) -> int:
    roman_map = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
    normalized = value.strip(" ().[]").upper()
    result = 0
    for index, char in enumerate(normalized):
        if index > 0 and roman_map[char] > roman_map[normalized[index - 1]]:
            result += roman_map[char] - 2 * roman_map[normalized[index - 1]]
        else:
            result += roman_map[char]
    return result


def alphabet_to_int(letter: str) -> int:
    normalized = letter.strip(" ().[]").lower()
    if not normalized.isalpha():
        return 1
    if len(normalized) == 1:
        return ord(normalized) - ord("a") + 1
    value = 0
    for char in normalized:
        value = value * 26 + (ord(char) - ord("a") + 1)
    return value


def detect_prefix(text: str):
    working = text or ""
    match = PREFIX_RE.match(working)
    if not match:
        return None, None, 0, "", "%1."

    full_match_str = match.group(0)
    prefix_len = len(full_match_str)
    leading_ws = match.group("leading") or ""
    opener = match.group("opener") or ""
    closer = match.group("closer") or "."
    value_str = match.group("value")
    clean_val = value_str.strip(" ().[]")
    template = f"{opener}%1{closer}"

    if clean_val.isalpha():
        if len(clean_val) > 2 and not is_roman(clean_val):
            return None, None, 0, "", "%1."
        if is_roman(clean_val):
            marker_type = "lowerRoman" if clean_val.islower() else "upperRoman"
        else:
            marker_type = "lowerLetter" if clean_val.islower() else "upperLetter"
        return marker_type, clean_val, prefix_len, leading_ws, template

    return "decimal", clean_val, prefix_len, leading_ws, template


def _text_indent_level(leading_ws: str) -> int:
    expanded_width = 0
    for char in leading_ws or "":
        expanded_width += 4 if char == "\t" else 1
    return min(8, expanded_width // 4)


def _style_list_level(paragraph) -> int:
    style_name = paragraph.style.name.strip() if paragraph.style and paragraph.style.name else ""
    match = STYLE_LEVEL_RE.match(style_name)
    if not match:
        return 0
    return max(0, min(8, int(match.group(1)) - 1))


def _paragraph_indent_level(paragraph) -> int:
    left_indent = paragraph.paragraph_format.left_indent
    if left_indent is None:
        return 0

    twips = getattr(left_indent, "twips", 0) or 0
    if twips < 720:
        return 0
    return max(0, min(8, round((twips - 720) / 720)))


def infer_list_level(paragraph, leading_ws: str) -> int:
    return max(
        _text_indent_level(leading_ws),
        _style_list_level(paragraph),
        _paragraph_indent_level(paragraph),
    )


def build_level_template(template: str, ilvl: int) -> str:
    return (template or "%1.").replace("%1", f"%{ilvl + 1}")


def get_or_create_abstract_num(doc, unique_id: int, fmt: str = "decimal", template: str = "%1."):
    num_part = doc.part.numbering_part
    if not num_part:
        doc.add_paragraph("", style="List Number")
        num_part = doc.part.numbering_part

    numbering_el = num_part.element
    for abs_num in numbering_el.findall(qn("w:abstractNum")):
        if abs_num.get(qn("w:abstractNumId")) == str(unique_id):
            return abs_num

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(unique_id))

    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level_type)

    for ilvl in range(9):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)

        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), build_level_template(template, ilvl))
        lvl.append(lvl_text)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        base_left = 1080 if "Letter" in fmt else 720
        ind.set(qn("w:left"), str(base_left + (720 * ilvl)))
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl.append(p_pr)

        abstract_num.append(lvl)

    numbering_el.insert(0, abstract_num)
    return abstract_num


def set_numpr_with_override(paragraph, doc, abstract_id: int, num_id: int, start_val: int, ilvl: int = 0) -> None:
    numbering_el = doc.part.numbering_part.element

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))

    abs_num_id = OxmlElement("w:abstractNumId")
    abs_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_num_id)

    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), str(ilvl))

    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start_val))
    lvl_override.append(start_override)

    num.append(lvl_override)
    numbering_el.append(num)

    p_pr = paragraph._element.get_or_add_pPr()
    _insert_pPr_child(p_pr, build_numPr(ilvl, num_id))


class ListProcessor:
    def __init__(self) -> None:
        self.active_lists: dict[int, dict[str, object]] = {}
        self.next_num_id = 5000
        self.abstract_registry: dict[tuple[str, str], int] = {}
        self.next_abstract_id = 9900

    def iter_paragraphs(self, doc, include_hf: bool = True):
        from docx.text.paragraph import Paragraph

        roots = [doc._element.body]
        if include_hf and doc.sections:
            seen_ids = set()
            section_attrs = [
                "header",
                "footer",
                "first_page_header",
                "first_page_footer",
                "even_page_header",
                "even_page_footer",
            ]
            for section in doc.sections:
                for attr in section_attrs:
                    try:
                        part = getattr(section, attr, None)
                        if part and id(part._element) not in seen_ids:
                            seen_ids.add(id(part._element))
                            roots.append(part._element)
                    except Exception:
                        continue
        for root in roots:
            for p_el in root.xpath(".//w:p"):
                yield Paragraph(p_el, doc)

    def _clean_text_layer(self, paragraph, prefix_len: int) -> None:
        chars_to_remove = prefix_len
        for run in paragraph.runs:
            if chars_to_remove <= 0:
                break

            if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
                continue

            text = run.text or ""
            if not text:
                continue

            if len(text) <= chars_to_remove:
                chars_to_remove -= len(text)
                run.text = ""
            else:
                run.text = text[chars_to_remove:]
                chars_to_remove = 0

    def _drop_deeper_states(self, ilvl: int) -> None:
        self.active_lists = {
            level: state for level, state in self.active_lists.items() if level <= ilvl
        }

    def _clear_states_from_level(self, ilvl: int) -> None:
        self.active_lists = {
            level: state for level, state in self.active_lists.items() if level < ilvl
        }

    def _process_document_impl(self, doc, include_hf: bool = True) -> ListFixStats:
        stats = ListFixStats()
        preserver = DocumentPreserver(doc)
        preserver.snapshot(include_hf=include_hf)

        paras = list(self.iter_paragraphs(doc, include_hf))
        cache = [(paragraph, "".join(run.text for run in paragraph.runs if run.text)) for paragraph in paras]

        bullet_num_id = self._ensure_bullet_def(doc)

        for paragraph, full_text in cache:
            if not full_text.strip() or has_numpr(paragraph):
                if not full_text.strip():
                    self.active_lists.clear()
                continue

            bullet_match = BULLET_RE.match(full_text)
            if bullet_match:
                ilvl = infer_list_level(paragraph, bullet_match.group("leading"))
                self._clear_states_from_level(ilvl)
                self._apply_list_meta(paragraph, ilvl, bullet_num_id)
                self._clean_text_layer(paragraph, len(bullet_match.group(0)))
                clear_paragraph_indents(paragraph)
                stats.unordered_tagged += 1
                continue

            marker_type, value_str, prefix_len, leading_ws, template = detect_prefix(full_text)
            if marker_type:
                ilvl = infer_list_level(paragraph, leading_ws)
                seq_val = self._to_int(value_str, marker_type)
                abs_id = self._get_abstract_id(doc, marker_type, template)

                self._drop_deeper_states(ilvl)
                state = self.active_lists.get(ilvl)
                is_new = (
                    not state
                    or state["m_type"] != marker_type
                    or state["template"] != template
                    or seq_val == 1
                    or seq_val != state["last_val"] + 1
                )

                if is_new:
                    self.next_num_id += 1
                    curr_num_id = self.next_num_id
                    set_numpr_with_override(paragraph, doc, abs_id, curr_num_id, seq_val, ilvl=ilvl)
                    self.active_lists[ilvl] = {
                        "id": curr_num_id,
                        "last_val": seq_val,
                        "m_type": marker_type,
                        "template": template,
                    }
                else:
                    self._apply_list_meta(paragraph, ilvl, int(state["id"]))
                    self.active_lists[ilvl]["last_val"] = seq_val

                self._clean_text_layer(paragraph, prefix_len)
                clear_paragraph_indents(paragraph)
                stats.ordered_tagged += 1
                continue

            self.active_lists.clear()

        stats.bullet_templates_standardized = self.standardize_bullet_symbols(doc)
        preserver.apply()
        return stats

    def process_document(self, input_path: str | os.PathLike[str], output_path: str | os.PathLike[str], include_hf: bool = True) -> ListFixStats:
        doc = Document(input_path)
        stats = self._process_document_impl(doc, include_hf)

        out_dir = os.path.dirname(os.path.abspath(output_path))
        if out_dir and not os.path.exists(out_dir):
            os.makedirs(out_dir, exist_ok=True)

        doc.save(output_path)
        return stats

    def process_document_bytes(self, input_bytes: bytes, include_hf: bool = True) -> tuple[bytes, ListFixStats]:
        doc = Document(BytesIO(input_bytes))
        stats = self._process_document_impl(doc, include_hf)

        output_bytes_io = BytesIO()
        doc.save(output_bytes_io)
        output_bytes_io.seek(0)
        return output_bytes_io.getvalue(), stats

    def _to_int(self, value_str: str, marker_type: str) -> int:
        if marker_type == "decimal":
            return int(value_str)
        if "Roman" in marker_type:
            return roman_to_int(value_str)
        return alphabet_to_int(value_str)

    def _apply_list_meta(self, paragraph, ilvl: int, num_id: int) -> None:
        p_pr = paragraph._element.get_or_add_pPr()
        _insert_pPr_child(p_pr, build_numPr(ilvl, num_id))

    def _ensure_bullet_def(self, doc, abstract_id: int = 9000, num_id: int = 9001) -> int:
        num_part = doc.part.numbering_part
        if not num_part:
            doc.add_paragraph("", style="List Bullet")
            num_part = doc.part.numbering_part

        numbering_el = num_part.element
        has_abs = any(
            abs_num.get(qn("w:abstractNumId")) == str(abstract_id)
            for abs_num in numbering_el.findall(qn("w:abstractNum"))
        )

        if not has_abs:
            abs_num = OxmlElement("w:abstractNum")
            abs_num.set(qn("w:abstractNumId"), str(abstract_id))

            multi_level_type = OxmlElement("w:multiLevelType")
            multi_level_type.set(qn("w:val"), "multilevel")
            abs_num.append(multi_level_type)

            for ilvl in range(9):
                symbol = BULLET_SYMBOLS[ilvl % len(BULLET_SYMBOLS)]
                lvl = OxmlElement("w:lvl")
                lvl.set(qn("w:ilvl"), str(ilvl))

                start = OxmlElement("w:start")
                start.set(qn("w:val"), "1")
                lvl.append(start)

                num_fmt = OxmlElement("w:numFmt")
                num_fmt.set(qn("w:val"), "bullet")
                lvl.append(num_fmt)

                lvl_text = OxmlElement("w:lvlText")
                lvl_text.set(qn("w:val"), symbol)
                lvl.append(lvl_text)

                lvl_jc = OxmlElement("w:lvlJc")
                lvl_jc.set(qn("w:val"), "left")
                lvl.append(lvl_jc)

                p_pr = OxmlElement("w:pPr")
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), str(720 * (ilvl + 1)))
                ind.set(qn("w:hanging"), "360")
                p_pr.append(ind)
                lvl.append(p_pr)

                abs_num.append(lvl)

            numbering_el.insert(0, abs_num)

        has_num = any(num.get(qn("w:numId")) == str(num_id) for num in numbering_el.findall(qn("w:num")))

        if not has_num:
            num = OxmlElement("w:num")
            num.set(qn("w:numId"), str(num_id))
            abs_num_id = OxmlElement("w:abstractNumId")
            abs_num_id.set(qn("w:val"), str(abstract_id))
            num.append(abs_num_id)
            numbering_el.append(num)

        return num_id

    def _get_abstract_id(self, doc, fmt: str, template: str) -> int:
        key = (fmt, template)
        if key not in self.abstract_registry:
            self.next_abstract_id += 1
            get_or_create_abstract_num(doc, self.next_abstract_id, fmt, template)
            self.abstract_registry[key] = self.next_abstract_id
        return self.abstract_registry[key]

    def standardize_bullet_symbols(self, doc) -> int:
        updated = 0
        try:
            num_part = doc.part.numbering_part
            if not num_part:
                return 0

            numbering_el = num_part.element
            for lvl in oxpath(numbering_el, ".//w:lvl"):
                ilvl = int(lvl.get(qn("w:ilvl"), "0"))
                num_fmt_el = lvl.find(qn("w:numFmt"))
                lvl_text_el = lvl.find(qn("w:lvlText"))
                if lvl_text_el is None:
                    continue

                value = lvl_text_el.get(qn("w:val"), "")
                fmt = num_fmt_el.get(qn("w:val"), "") if num_fmt_el is not None else ""

                is_bullet = (
                    fmt == "bullet"
                    or any(ord(char) >= 0xF000 for char in value)
                    or any(char in value for char in "•\xb7")
                )
                if not is_bullet:
                    continue

                target_symbol = BULLET_SYMBOLS[ilvl % len(BULLET_SYMBOLS)]
                lvl_text_el.set(qn("w:val"), target_symbol)

                if num_fmt_el is None:
                    num_fmt_el = OxmlElement("w:numFmt")
                    lvl.append(num_fmt_el)
                num_fmt_el.set(qn("w:val"), "bullet")

                for sym in oxpath(lvl, ".//w:sym"):
                    sym.getparent().remove(sym)

                r_pr = lvl.find(qn("w:rPr"))
                if r_pr is None:
                    r_pr = OxmlElement("w:rPr")
                    lvl.append(r_pr)

                r_fonts = r_pr.find(qn("w:rFonts"))
                if r_fonts is None:
                    r_fonts = OxmlElement("w:rFonts")
                    r_pr.insert(0, r_fonts)

                r_fonts.set(qn("w:ascii"), "Calibri")
                r_fonts.set(qn("w:hAnsi"), "Calibri")
                r_fonts.set(qn("w:eastAsia"), "Calibri")
                r_fonts.set(qn("w:cs"), "Calibri")
                r_fonts.set(qn("w:hint"), "default")
                updated += 1
        except Exception as exc:
            print(f"Warning: Bullet standardization skipped: {exc}")

        return updated


def build_list_correction_output_name(filename: str) -> str:
    path = Path(filename or "document.docx")
    return f"{path.stem}_list_corrected{path.suffix or '.docx'}"


def main() -> None:
    parser = argparse.ArgumentParser(description="Tag semantic lists in a DOCX and save a normalized copy.")
    parser.add_argument("input_docx", help="Source DOCX file path")
    parser.add_argument("output_docx", help="Destination DOCX file path")
    parser.add_argument(
        "--include-header-footer",
        dest="include_hf",
        action="store_true",
        default=False,
        help="Also process paragraph content in headers and footers",
    )
    parser.add_argument(
        "--exclude-header-footer",
        dest="include_hf",
        action="store_false",
        help="Skip headers and footers",
    )
    args = parser.parse_args()

    processor = ListProcessor()
    stats = processor.process_document(args.input_docx, args.output_docx, include_hf=args.include_hf)

    print(f"Processed: {args.input_docx}\n")
    print(f"Written:   {args.output_docx}\n")
    print(f"Ordered lists tagged:   {stats.ordered_tagged}\n")
    print(f"Unordered lists tagged: {stats.unordered_tagged}\n")
    print(f"Bullet templates normalized: {stats.bullet_templates_standardized}\n")


if __name__ == "__main__":
    main()
